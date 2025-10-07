# D:\ArsaWebDeveloper\Project\Vexta\routers\pdf.py

from fastapi import APIRouter, UploadFile, File, Depends, Request, HTTPException
from sqlalchemy.orm import Session
import os
from pydantic import BaseModel
import tempfile
from utils.pdf_utils import extract_text_from_pdf, chunk_text
from services.rag import store_chunks, search_similar_chunks, generate_answer
from database import get_db
from openai import OpenAI
from models import Message,PDFChunk
from datetime import datetime
from fastapi import APIRouter
from fastapi import APIRouter
from fastapi.responses import FileResponse
from pathlib import Path
from docx import Document
from auth.deps import get_current_user
from auth.models import User 
from fastapi import Depends
from fastapi.responses import FileResponse
import uuid
import fitz  # PyMuPDF



client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))




router = APIRouter()

MAX_SIZE_MB = 50

@router.post("/upload")
async def upload_pdf(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user) 
):
    try:
        contents = await file.read()

        if len(contents) > MAX_SIZE_MB * 1024 * 1024:
            raise HTTPException(status_code=400, detail="حجم فایل بیشتر از 50MB مجاز نیست")

        filename = file.filename
        if not filename or not filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail="فقط فایل PDF مجاز است")

        # 📂 ساخت پوشه مخصوص کاربر
        user_email = str(current_user.email)
        user_folder = os.path.join("static", "pdf", user_email)
        os.makedirs(user_folder, exist_ok=True)

        save_path = os.path.join(user_folder, filename)

        # 💾 ذخیره فایل
        with open(save_path, "wb") as f:
            f.write(contents)

        # 📄 پردازش و ذخیره چانک‌ها
        text = extract_text_from_pdf(save_path)
        chunks = chunk_text(text)

        user_id = int(current_user.id)
        store_chunks(chunks, filename, db, user_id)

        return {
            "message": "فایل با موفقیت آپلود و پردازش شد",
            "chunks": len(chunks),
            "filename": filename
        }

    except Exception as e:
        traceback.print_exc()   # 🛠 خطا رو کامل در لاگ چاپ کن
        raise HTTPException(status_code=500, detail=f"خطا در پردازش فایل: {str(e)}")
    

# --- فقط همین تابع را جایگزین کن ---
@router.get("/pdf-to-word/{filename}")
def pdf_to_word(
    filename: str,
    current_user=Depends(get_current_user)
):
    from pathlib import Path
    import fitz  # PyMuPDF
    from docx import Document
    from fastapi import HTTPException

    # مسیر PDF
    pdf_path = Path("static") / "pdf" / str(current_user.email) / filename

    if not pdf_path.exists():
        raise HTTPException(status_code=404, detail="فایل پیدا نشد")

    # 📄 استخراج متن از PDF
    all_text = ""
    with fitz.open(str(pdf_path)) as doc:
        for page in doc:
            all_text += page.get_text("text") + "\n\n"  # type: ignore

    # 📝 ساخت Word در پوشه اختصاصی کاربر
    out_dir = Path("converted") / str(current_user.email)
    out_dir.mkdir(parents=True, exist_ok=True)
    word_path = out_dir / (pdf_path.stem + ".docx")

    word_doc = Document()
    word_doc.add_paragraph(all_text if all_text.strip() else " ")
    word_doc.save(str(word_path))

    return FileResponse(
        path=str(word_path),
        filename=word_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# --- پایان: فقط همین تابع ---

@router.get("/pdf-to-image/{filename}/{page_number}")
def pdf_to_image(
    filename: str,
    page_number: int,
    current_user=Depends(get_current_user)
):
    pdf_path = Path("static") / "pdf" / str(current_user.email) / filename
    print("📂 مسیر فایل:", pdf_path)

    if not pdf_path.exists():
        raise HTTPException(status_code=404, detail="فایل پیدا نشد")

    doc = fitz.open(str(pdf_path))
    print("📄 تعداد صفحات PDF:", len(doc))

    if page_number < 1 or page_number > len(doc):
        doc.close()
        raise HTTPException(status_code=400, detail="شماره صفحه نامعتبر است")

    page = doc[page_number - 1]
    pix = page.get_pixmap(dpi=150)  # type: ignore
    print("🖼️ سایز تصویر:", pix.width, "x", pix.height)

    doc.close()

    output_dir = Path("converted") / str(current_user.email)
    output_dir.mkdir(parents=True, exist_ok=True)
    img_path = output_dir / f"{pdf_path.stem}_page{page_number}.jpg"
    pix.save(str(img_path))

    return FileResponse(
        path=str(img_path),
        filename=img_path.name,
        media_type="image/jpeg"
    )

@router.get("/translate/{filename}")
async def translate_pdf(
    filename: str, 
    lang: str = "fa", 
    current_user = Depends(get_current_user)
):
    user_email = str(current_user.email)
    pdf_path = os.path.join("static", "pdf", user_email, filename)

    if not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail="فایل پیدا نشد")

    # استخراج متن PDF
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text("text") + "\n" # type: ignore
    doc.close()

    # ترجمه با OpenAI
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a professional translator."},
            {"role": "user", "content": f"Translate the following text into {lang}:\n{text}"}
        ]
    )
    translated_text = response.choices[0].message.content

    # ذخیره فایل Word
    output_dir = os.path.join("static", "translated", user_email)
    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, f"{uuid.uuid4()}.docx")

    word_doc = Document()
    word_doc.add_paragraph(translated_text or "")

    word_doc.save(output_file)

    return FileResponse(output_file, filename=f"{filename}_translated.docx")




@router.get("/translate-text/{filename}")
async def translate_text(
    filename: str, 
    lang: str = "fa", 
    current_user = Depends(get_current_user)
):
    user_email = str(current_user.email)
    pdf_path = os.path.join("static", "pdf", user_email, filename)

    if not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail="فایل پیدا نشد")

    # استخراج متن
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text("text") for page in doc]) # type: ignore
    doc.close()

    # ترجمه با OpenAI
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a professional translator."},
            {"role": "user", "content": f"Translate the following text into {lang}:\n{text}"}
        ]
    )
    translated_text = response.choices[0].message.content or ""

    return {"translated_text": translated_text}




import traceback  # بالای فایل اضافه کن

@router.post("/ask")
async def ask_pdf(request: Request, db: Session = Depends(get_db)):
    try:
        data = await request.json()
        question = data.get("question")
        filename = data.get("filename")

        print("داده‌های دریافتی:", data)

        if not question or not filename:
            raise HTTPException(status_code=400, detail="سؤال یا نام فایل ارسال نشده است")

        chunks = search_similar_chunks(question, db, filename)
        answer = generate_answer(question, chunks)

        # ذخیره در جدول Message
        message = Message(
            filename=filename,
            user_message=question,
            bot_reply=answer,
            created_at=datetime.utcnow()
        )
        db.add(message)
        db.commit()

        return {"answer": answer}

    except Exception as e:
        traceback.print_exc()  # 🛠️ خطای کامل در ترمینال چاپ شود
        raise HTTPException(status_code=500, detail=f"خطا در پردازش سؤال: {str(e)}")


@router.get("/debug/chunks/{filename}")
def debug_chunks(filename: str, db: Session = Depends(get_db)):
    chunks = db.query(PDFChunk).filter(PDFChunk.filename == filename).all()
    return {"total_chunks": len(chunks), "chunks": [chunk.chunk[:100] for chunk in chunks]}


@router.get("/history/{filename}")
def get_chat_history(filename: str, db: Session = Depends(get_db)):
    messages = db.query(Message).filter(Message.filename == filename).order_by(Message.created_at.asc()).all()
    return [
        {
            "question": msg.user_message,
            "answer": msg.bot_reply,
            "timestamp": msg.created_at.isoformat()
        }
        for msg in messages
    ]


@router.get("/files")
def list_uploaded_files(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user)
):
    user_folder = os.path.join("static", "pdf", str(current_user.email))
    files = []
    if os.path.isdir(user_folder):
        files = [f for f in os.listdir(user_folder) if f.lower().endswith(".pdf")]
    return sorted(files)






@router.delete("/{filename}")
def delete_pdf(
    filename: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    user_email = getattr(current_user, "email", None)
    if not user_email:
        raise HTTPException(status_code=400, detail="ایمیل کاربر یافت نشد")

    # مسیر فایل اصلی PDF
    file_path = os.path.join("static", "pdf", str(user_email), filename)

    file_existed = os.path.exists(file_path)
    if file_existed:
        try:
            os.remove(file_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"خطا در حذف فایل PDF: {str(e)}")

    # 🗑 پاک کردن خروجی‌های مرتبط (Word, Images, Translations)
    base_name = Path(filename).stem  # مثلا report.pdf → report

    # پوشه خروجی‌های کاربر
    converted_dir = Path("converted") / str(user_email)
    translated_dir = Path("static") / "translated" / str(user_email)

    # پاک کردن فایل‌های Word/Image در converted
    if converted_dir.exists():
        for f in converted_dir.iterdir():
            if f.is_file() and f.stem.startswith(base_name):
                try:
                    f.unlink()
                except Exception as e:
                    print(f"❌ خطا در حذف {f}: {e}")

    # پاک کردن فایل‌های ترجمه‌شده در translated
    if translated_dir.exists():
        for f in translated_dir.iterdir():
            if f.is_file() and f.name.startswith(base_name):
                try:
                    f.unlink()
                except Exception as e:
                    print(f"❌ خطا در حذف {f}: {e}")

    # 🗑 پاک کردن رکوردهای DB مربوط به فایل
    chunk_deleted = db.query(PDFChunk).filter(
        PDFChunk.filename == filename,
        PDFChunk.user_id == current_user.id
    ).delete(synchronize_session=False)

    message_deleted = db.query(Message).filter(
        Message.filename == filename,
        Message.user_id == current_user.id
    ).delete(synchronize_session=False)

    db.commit()

    if not file_existed and chunk_deleted == 0 and message_deleted == 0:
        raise HTTPException(status_code=404, detail="فایلی با این نام برای شما پیدا نشد")

    return {
        "message": "حذف انجام شد",
        "file_removed": file_existed,
        "chunks_removed": chunk_deleted,
        "messages_removed": message_deleted
    }


@router.get("/view/{filename}")
def view_pdf(
    filename: str,
    current_user=Depends(get_current_user)
):
    file_path = os.path.join("static", "pdf", str(current_user.email), filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="فایل پیدا نشد")

    # inline نمایش
    return FileResponse(
        path=file_path,
        media_type="application/pdf",
        filename=filename,
        headers={"Content-Disposition": f'inline; filename="{filename}"'}
    )


@router.get("/summarize/{filename}")
async def summarize_pdf(
    filename: str,
    lines: int = 0,              # مثلا 10 خط
    mode: str = "paragraph",     # "paragraph" یا "short"
    lang: str = "fa",
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user)
):
    user_email = str(current_user.email)
    file_path = os.path.join("static", "pdf", user_email, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="فایل پیدا نشد")

    # 📄 استخراج متن کامل PDF
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"خطا در خواندن PDF: {e}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="متنی در PDF پیدا نشد")

    # ✂️ تقسیم متن به چانک‌های کوچک‌تر (مثلا 1500 کاراکتر)
    chunks = chunk_text(text, max_tokens=1500)

    summaries = []
    for idx, chunk in enumerate(chunks, start=1):
        # آماده‌سازی پرامپت برای هر چانک
        prompt = f"این متن را به {lang} خلاصه کن.\n\n"
        if lines > 0:
            prompt += f"خلاصه باید حدود {lines} خط باشد.\n\n"
        elif mode == "paragraph":
            prompt += "خلاصه باید در یک پاراگراف کوتاه باشد.\n\n"
        elif mode == "short":
            prompt += "خلاصه باید خیلی کوتاه (چند جمله) باشد.\n\n"

        prompt += f"متن:\n{chunk}"

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
            )
            content = response.choices[0].message.content
            summary = content.strip() if content else ""
            summaries.append(summary)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"خطا در خلاصه‌سازی چانک {idx}: {e}")

    # 📝 ترکیب خلاصه‌های همه چانک‌ها در یک خروجی نهایی
    final_prompt = f"خلاصه زیر را ترکیب کن و یک جمع‌بندی نهایی به {lang} بده:\n\n" + "\n\n".join(summaries)




    try:
        final_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": final_prompt}],
            temperature=0.7,
        )
        content = final_response.choices[0].message.content
        final_summary = content.strip() if content else "" 
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"خطا در ترکیب خلاصه‌ها: {e}")

    return {"summary": final_summary, "chunks_summarized": len(chunks)}






